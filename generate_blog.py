"""
[*] AI+패스웨이 블로그 DOCX 자동생성
================================
Supabase에서 분석 결과를 읽어 블로그 DOCX를 자동 생성

사용법:
  python generate_blog.py                          -> 최신 날짜, TOP3 자동
  python generate_blog.py 20260320                 -> 날짜 지정
  python generate_blog.py 20260320 "SNT에너지,태웅,네패스아크"  -> TOP 직접 지정
  python generate_blog.py 20260320 "SNT에너지,태웅,네패스아크" "비에이치"  -> TOP + 추천주

필요: pip install python-docx requests python-dotenv
"""

import os
import sys
import time
import requests
from datetime import datetime, timedelta
from dotenv import load_dotenv

load_dotenv()

SUPABASE_URL = os.getenv("SUPABASE_URL", "").rstrip('/')
SUPABASE_KEY = os.getenv("SUPABASE_KEY", "")

HEADERS = {
    "apikey": SUPABASE_KEY,
    "Authorization": f"Bearer {SUPABASE_KEY}",
}


def db_read(table, params=""):
    all_data = []
    offset = 0
    limit = 1000
    while True:
        url = f"{SUPABASE_URL}/rest/v1/{table}?{params}&limit={limit}&offset={offset}" if params else f"{SUPABASE_URL}/rest/v1/{table}?limit={limit}&offset={offset}"
        resp = requests.get(url, headers=HEADERS)
        if resp.status_code != 200:
            break
        data = resp.json()
        if not data:
            break
        all_data.extend(data)
        if len(data) < limit:
            break
        offset += limit
    return all_data


def get_sector(code, sector_map):
    return sector_map.get(code, '기타')


def generate_blog(target_date, top_names=None, recommend_name=None):
    try:
        from docx import Document
        from docx.shared import Pt, Inches, Cm, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
    except ImportError:
        print("[X] python-docx 필요: pip install python-docx --break-system-packages")
        return

    date_str = datetime.strptime(target_date, "%Y%m%d").strftime("%Y-%m-%d")
    date_display = datetime.strptime(target_date, "%Y%m%d").strftime("%Y년 %m월 %d일")
    weekdays = ['월', '화', '수', '목', '금', '토', '일']
    weekday = weekdays[datetime.strptime(target_date, "%Y%m%d").weekday()]

    print(f"[D] 블로그 생성: {date_str} ({weekday})")

    # 데이터 로드
    scores = db_read("analysis_scores", f"date=eq.{date_str}&order=final_score.desc")
    if not scores:
        print("[X] 분석 데이터가 없습니다.")
        return

    # 섹터맵 로드
    sector_data = db_read("sector_map")
    sector_map = {r['stock_code']: r['sector'] for r in sector_data}

    five_all = [s for s in scores if s.get('n_buyers') == 5]
    d_strategy = [s for s in scores if s.get('combo') and '외' in s['combo'] and '연' in s['combo'] and '사' in s['combo']]

    # TOP 결정
    if top_names:
        name_list = [n.strip() for n in top_names.split(',')]
        top_stocks = []
        for name in name_list:
            found = [s for s in scores if s['stock_name'] == name]
            if found:
                top_stocks.append(found[0])
        print(f"  TOP{len(top_stocks)}: {', '.join(s['stock_name'] for s in top_stocks)}")
    else:
        no_conflict = [s for s in scores if not s.get('conflicts')]
        top_stocks = no_conflict[:3] if len(no_conflict) >= 3 else scores[:3]
        print(f"  TOP3 (자동): {', '.join(s['stock_name'] for s in top_stocks)}")

    # TOP3 히스토리 자동 저장 (Supabase)
    save_top3_history(target_date, top_stocks)

    # 추천주
    recommend = None
    if recommend_name:
        found = [s for s in scores if s['stock_name'] == recommend_name.strip()]
        if found:
            recommend = found[0]
            print(f"  추천주: {recommend['stock_name']}")

    # 전일 TOP3 성과
    curr_date = datetime.strptime(date_str, "%Y-%m-%d")
    prev_date = curr_date - timedelta(days=1)
    while prev_date.weekday() >= 5:
        prev_date -= timedelta(days=1)
    prev_date_str = prev_date.strftime("%Y-%m-%d")
    prev_scores = db_read("analysis_scores", f"date=eq.{prev_date_str}&order=final_score.desc&limit=10")
    prev_top3 = [r for r in prev_scores if not r.get('conflicts', '')][:3] if prev_scores else []

    today_market = db_read("daily_market", f"date=eq.{date_str}")
    mkt = {r['stock_code']: r for r in today_market}

    # ========== DOCX 생성 ==========
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style.font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    def add_title(text, size=18, color=RGBColor(0x2F, 0x54, 0x96), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
        p = doc.add_paragraph()
        p.alignment = align
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = color
        run.font.bold = bold
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        return p

    def add_body(text, size=11):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        return p

    def set_cell(cell, text, size=9, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER, bg=None):
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = align
        run = p.add_run(str(text))
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.name = '맑은 고딕'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
        if color:
            run.font.color.rgb = color
        if bg:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): bg, qn('w:val'): 'clear'
            })
            shading.append(shading_elm)

    # 제목
    add_title("[G] AI 수급분석 리포트", size=20)
    add_title(f"{date_display} ({weekday}) | AI+패스웨이", size=12, color=RGBColor(0x66, 0x66, 0x66), bold=False)
    doc.add_paragraph()

    # 시장 총평
    add_title("# 시장 총평", size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_body(f"5주체 전원매수 {len(five_all)}종목, 3주체^ {len([s for s in scores if s.get('n_buyers',0)>=3])}종목, D전략(외+연+사) {len(d_strategy)}종목이 포착되었습니다.")
    if five_all:
        sectors = {}
        for s in five_all:
            sec = get_sector(s['stock_code'], sector_map)[:10]
            sectors[sec] = sectors.get(sec, 0) + 1
        top_sectors = sorted(sectors.items(), key=lambda x: x[1], reverse=True)[:3]
        sector_text = ', '.join(f"{s}({c}종목)" for s, c in top_sectors)
        add_body(f"주요 업종: {sector_text}")
    doc.add_paragraph()

    # TOP 종목
    add_title(f"# 오늘의 TOP{len(top_stocks)} 종목", size=14, color=RGBColor(0xC0, 0x00, 0x00), align=WD_ALIGN_PARAGRAPH.LEFT)
    cols = [1.0, 2.5, 3.0, 1.0, 1.2, 2.5]
    table = doc.add_table(rows=1 + len(top_stocks), cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['순위', '종목명', '주요제품', '주체수', '점수', '조합']
    for i, h in enumerate(headers):
        set_cell(table.rows[0].cells[i], h, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), bg='2F5496')
    for idx, stock in enumerate(top_stocks):
        row = table.rows[idx + 1]
        bg = 'FFF2CC' if idx == 0 else ('E2EFDA' if idx == 1 else 'D6E4F0')
        set_cell(row.cells[0], f"{idx+1}위", bold=True, bg=bg)
        set_cell(row.cells[1], stock['stock_name'], bold=True, bg=bg, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[2], get_sector(stock['stock_code'], sector_map)[:30], bg=bg, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell(row.cells[3], str(stock.get('n_buyers', '')), bg=bg)
        set_cell(row.cells[4], f"{stock.get('final_score', 0):.1f}", bold=True, bg=bg, color=RGBColor(0xC0, 0x00, 0x00))
        set_cell(row.cells[5], stock.get('combo', ''), bg=bg)
    doc.add_paragraph()

    # TOP 종목 상세
    for idx, stock in enumerate(top_stocks):
        sector = get_sector(stock['stock_code'], sector_map)
        chg = stock.get('change_pct', 0)
        add_title(f"[>] {idx+1}위: {stock['stock_name']} ({stock.get('final_score',0):.1f}점)", size=12, align=WD_ALIGN_PARAGRAPH.LEFT)
        add_body(f"주요제품: {sector}")
        add_body(f"{stock.get('n_buyers',0)}주체 매수 ({stock.get('combo','')}), 당일 {chg:+.2f}%")
        doc.add_paragraph()

    # 추천주
    if recommend:
        add_title(f"# 추천주: {recommend['stock_name']}", size=14, color=RGBColor(0x70, 0x30, 0xA0), align=WD_ALIGN_PARAGRAPH.LEFT)
        rec_sector = get_sector(recommend['stock_code'], sector_map)
        add_body(f"주요제품: {rec_sector}")
        add_body(f"{recommend.get('n_buyers',0)}주체 매수 ({recommend.get('combo','')}), 점수 {recommend.get('final_score',0):.1f}, 당일 {recommend.get('change_pct',0):+.2f}%")
        doc.add_paragraph()

    # 5주체 전원매수
    if five_all:
        add_title(f"# 5주체 전원매수 ({len(five_all)}종목)", size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
        t2 = doc.add_table(rows=1 + len(five_all), cols=4)
        t2.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['종목명', '주요제품', '점수', '등락률']):
            set_cell(t2.rows[0].cells[i], h, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), bg='2F5496')
        for idx, stock in enumerate(five_all):
            row = t2.rows[idx + 1]
            bg = 'E2EFDA' if idx % 2 == 0 else None
            set_cell(row.cells[0], stock['stock_name'], align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg)
            set_cell(row.cells[1], get_sector(stock['stock_code'], sector_map)[:30], align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg)
            set_cell(row.cells[2], f"{stock.get('final_score',0):.1f}", bg=bg)
            chg = stock.get('change_pct', 0)
            c = RGBColor(0xFF, 0x00, 0x00) if chg > 0 else RGBColor(0x00, 0x00, 0xFF)
            set_cell(row.cells[3], f"{chg:+.2f}%", color=c, bg=bg)
        doc.add_paragraph()

    # 전일 TOP3 성과
    if prev_top3:
        add_title("# 전일 TOP3 성과", size=14, align=WD_ALIGN_PARAGRAPH.LEFT)
        t3 = doc.add_table(rows=1 + len(prev_top3) + 1, cols=3)
        t3.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, h in enumerate(['종목명', '주요제품', '익일등락률']):
            set_cell(t3.rows[0].cells[i], h, size=9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), bg='2F5496')
        chg_vals = []
        for idx, stock in enumerate(prev_top3):
            row = t3.rows[idx + 1]
            code = stock['stock_code']
            today_chg = mkt.get(code, {}).get('change_pct', 0) or 0
            chg_vals.append(today_chg)
            set_cell(row.cells[0], stock['stock_name'], align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell(row.cells[1], get_sector(code, sector_map)[:30], align=WD_ALIGN_PARAGRAPH.LEFT)
            c = RGBColor(0xFF, 0x00, 0x00) if today_chg > 0 else RGBColor(0x00, 0x00, 0xFF)
            set_cell(row.cells[2], f"{today_chg:+.2f}%", color=c)
        avg_row = t3.rows[len(prev_top3) + 1]
        avg_chg = sum(chg_vals) / len(chg_vals) if chg_vals else 0
        set_cell(avg_row.cells[0], '평균', bold=True, bg='D6E4F0')
        set_cell(avg_row.cells[1], '', bg='D6E4F0')
        c = RGBColor(0xFF, 0x00, 0x00) if avg_chg > 0 else RGBColor(0x00, 0x00, 0xFF)
        set_cell(avg_row.cells[2], f"{avg_chg:+.2f}%", bold=True, color=c, bg='D6E4F0')
        doc.add_paragraph()

    # 마무리
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('"남들이 차트를 볼 때, 저는 수급을 봅니다."')
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run.font.name = '맑은 고딕'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('"남들이 공포에 팔 때, 기관이 어디를 사는지 봅니다."')
    run2.font.size = Pt(12)
    run2.font.bold = True
    run2.font.color.rgb = RGBColor(0x2F, 0x54, 0x96)
    run2.font.name = '맑은 고딕'

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run('AI+패스웨이 | 본 리포트는 투자 참고용이며 투자 판단의 책임은 본인에게 있습니다.')
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    # 저장
    date_short = datetime.strptime(target_date, "%Y%m%d").strftime("%y%m%d")
    filename = f"Claude_블로그_수급분석_{date_short}.docx"
    filepath = os.path.join(".", filename)
    doc.save(filepath)
    print(f"\n[OK] 블로그 생성 완료: {filepath}")
    return filepath

def save_top3_history(target_date, top_stocks):
    """TOP3 확정 시 top3_history 테이블에 자동 저장"""
    date_str = f"{target_date[:4]}-{target_date[4:6]}-{target_date[6:8]}"
    headers = {"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}", "Content-Type": "application/json", "Prefer": "resolution=merge-duplicates"}
    
    # 해당 날짜의 daily_market에서 종가 가져오기
    price_map = {}
    for s in top_stocks:
        code = s.get('stock_code', '')
        url = f"{SUPABASE_URL}/rest/v1/daily_market?stock_code=eq.{code}&date=eq.{date_str}&select=close_price"
        resp = requests.get(url, headers={"apikey": SUPABASE_KEY, "Authorization": f"Bearer {SUPABASE_KEY}"})
        if resp.status_code == 200 and resp.json():
            price_map[code] = resp.json()[0]['close_price']
    
    # 20영업일 후 만료일 계산
    expire = datetime.strptime(target_date, "%Y%m%d")
    biz_days = 0
    while biz_days < 20:
        expire += timedelta(days=1)
        if expire.weekday() < 5:
            biz_days += 1
    expires_date = expire.strftime("%Y-%m-%d")
    
    rows = []
    for i, s in enumerate(top_stocks):
        code = s.get('stock_code', '')
        base_price = price_map.get(code, 0)
        rows.append({
            "date": date_str,
            "rank": i + 1,
            "stock_code": code,
            "stock_name": s.get('stock_name', ''),
            "base_price": base_price,
            "score": s.get('final_score', 0),
            "combo": s.get('combo', ''),
            "sector": s.get('sector', ''),
            "selected_by": "shawn",
            "expires_date": expires_date
        })
    
    url = f"{SUPABASE_URL}/rest/v1/top3_history"
    resp = requests.post(url, headers=headers, json=rows)
    if resp.status_code in [200, 201]:
        print(f"  [OK] top3_history 저장 완료 ({len(rows)}종목, 만료일: {expires_date})")
    else:
        print(f"  [W] top3_history 저장 실패: {resp.status_code} {resp.text[:100]}")

def main():
    target_date = sys.argv[1] if len(sys.argv) > 1 else datetime.now().strftime("%Y%m%d")
    top_names = sys.argv[2] if len(sys.argv) > 2 else None
    recommend = sys.argv[3] if len(sys.argv) > 3 else None

    print("=" * 50)
    print("[*] AI+패스웨이 블로그 자동생성")
    print(f"   날짜: {target_date}")
    print("=" * 50)

    if not SUPABASE_URL or not SUPABASE_KEY:
        print("[X] .env 파일에 SUPABASE_URL, SUPABASE_KEY를 설정하세요.")
        return

    generate_blog(target_date, top_names, recommend)


if __name__ == "__main__":
    main()
